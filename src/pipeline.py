"""
Pipeline CLI : audio → transcription → structuration → export Word
Usage :
  python pipeline.py                        # mode mock complet
  python pipeline.py --audio fichier.mp3    # avec vrai fichier audio (clé Gladia requise)
  python pipeline.py --transcription t.txt  # depuis une transcription existante
"""

import argparse
import base64
import io
import json
import os
import sys
from pathlib import Path
from datetime import date, datetime, timezone

# Charge .env si présent (python-dotenv optionnel)
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent.parent / ".env")
except ImportError:
    pass

# --------------------------------------------------------------------------- #
# 1. TRANSCRIPTION (mock ou Whisper self-hosted)
# --------------------------------------------------------------------------- #

MOCK_TRANSCRIPTION = """
Alors bonjour à tous, on est le 28 mai 2025, on se retrouve sur le chantier de la maison Dupont,
au 12 rue des Lilas à Lyon. Sont présents : moi-même Jean-Marc Aubert, architecte du cabinet Aubert
et Associés, Michel Renard de chez Bâticorp pour le gros œuvre, Sophie Leclaire de chez ElecPro
pour l'électricité, et Thomas Blanc de Menuiserie Blanche. Monsieur Fabre de la plomberie est absent,
il nous a prévenus ce matin.

Alors on commence par le gros œuvre. Michel, où en est-on sur le dallage du garage ?
Michel dit que le coulage est prévu jeudi prochain, le premier juin. Il attend la confirmation
météo mais ça devrait être bon. Par contre il signale un problème sur le mur de refend au niveau
du salon, il y a une fissure d'environ deux millimètres qui est apparue, il faut qu'on regarde ça
ensemble. On décide de faire un constat la semaine prochaine avec un bureau de contrôle.
Jean-Marc note que Michel doit envoyer le PV de résistance du béton avant vendredi.

Pour la charpente, elle est posée, RAS, nickel. Pas de remarque.

On passe à l'électricité. Sophie explique que le passage des gaines dans les cloisons du premier
étage est terminé à quatre-vingts pour cent. Il reste la chambre parentale et la salle de bain.
Prévu pour fin de semaine prochaine. Elle demande à avoir les plans de cuisine définitifs pour
pouvoir positionner les prises, elle attend ça de l'architecte. Jean-Marc dit qu'il envoie ça
ce soir par email.

Menuiseries extérieures avec Thomas. Les fenêtres du rez-de-chaussée sont posées, mais Thomas
signale que la fenêtre de la cuisine a un problème de joint, elle ferme mal. Il faut une
intervention du fabricant, Thomas prend contact cette semaine. Les fenêtres du premier étage
ne sont pas encore livrées, retard fournisseur, nouveau délai annoncé au quinze juin.

Questions diverses : Jean-Marc rappelle que l'accès au chantier doit être sécurisé, il manque
un cadenas sur le portail. Michel s'en charge demain.

Prochaine réunion le mardi dix juin à quatorze heures sur le chantier.
Diffusion du CR à tous les présents plus Monsieur Fabre et le maître d'ouvrage Monsieur Dupont.
"""


def transcribe_audio(audio_path: str) -> str:
    """Transcription locale via faster-whisper (self-hosted, gratuit)."""
    try:
        from faster_whisper import WhisperModel
    except ImportError:
        print("[ERREUR] pip install faster-whisper")
        sys.exit(1)

    model_size = os.environ.get("WHISPER_MODEL", "medium")
    device = os.environ.get("WHISPER_DEVICE", "cpu")
    compute_type = "int8" if device == "cpu" else "float16"

    print(f"Chargement du modèle Whisper {model_size} ({device})...")
    model = WhisperModel(model_size, device=device, compute_type=compute_type)

    print(f"Transcription de {audio_path} en cours...")
    segments, info = model.transcribe(
        audio_path,
        language="fr",
        beam_size=5,
        vad_filter=True,
        initial_prompt=(
            "Réunion de chantier. Participants : architecte, maître d'œuvre, entreprises. "
            "Termes : lot, gros œuvre, charpente, menuiserie, plomberie, DTU, réserve, levée de réserve."
        ),
    )

    print(f"Langue détectée : {info.language} (probabilité {info.language_probability:.0%})")
    lines = [segment.text.strip() for segment in segments]
    return "\n".join(lines)


def get_transcription(args) -> str:
    if args.transcription:
        return Path(args.transcription).read_text(encoding="utf-8")
    if args.audio:
        return transcribe_audio(args.audio)
    print("[MOCK] Utilisation de la transcription de démonstration.")
    return MOCK_TRANSCRIPTION








# --------------------------------------------------------------------------- #
# 2. STRUCTURATION LLM (mock ou Mistral réel)
# --------------------------------------------------------------------------- #

SYSTEM_PROMPT = """Tu es un assistant spécialisé dans la rédaction de comptes rendus de chantier pour architectes français.

À partir de la transcription d'une réunion de chantier, tu dois produire un compte rendu structuré au format JSON.

Règles :
- Identifie tous les intervenants mentionnés (nom, qualité, entreprise si mentionné)
- Classe chaque point soulevé par lot (gros œuvre, charpente, électricité, plomberie, menuiserie, etc.)
- Pour chaque point, identifie : la description, la décision prise (si applicable), l'action à mener, le responsable, le délai
- Identifie la date et le lieu de la prochaine réunion si mentionnés
- En cas d'ambiguïté, garde les deux versions possibles avec une note [À VÉRIFIER]
- Ne jamais inventer d'informations non présentes dans la transcription
- La transcription peut contenir des erreurs de reconnaissance vocale typiques du français : corrige-les par le contexte. Exemples fréquents : "l'eau" → "lot", "Mohamed" ou un prénom isolé → probablement "moi-même" (l'architecte qui parle), "l'eau numéro" → "lot numéro", "lever" → "livrer" selon contexte, "réserve" conservé tel quel car terme métier.
- Pour la numérotation des lots : si elle n'est pas explicitement mentionnée dans la transcription, numéroter séquentiellement 01, 02, 03, etc. dans l'ordre d'apparition. Ne jamais inventer un numéro de lot.
- Pour chaque action identifiée, attribuer le responsable à la personne qui s'engage à la faire, pas à celle qui la reçoit. Si l'architecte dit "je lui envoie" ou "j'envoie", le responsable est l'architecte. Si une entreprise "s'engage" ou "s'en charge", le responsable est cette entreprise.
- Pour les délais : convertir "vendredi prochain" en date absolue en se basant sur la date de réunion fournie. Vendredi suivant le 2026-06-01 = 2026-06-05.
- Les descriptions de points doivent être précises et techniques : inclure les dimensions chiffrées, les références de matériaux, les localisations exactes mentionnées dans la transcription. Éviter les formulations vagues.
- Si un point mentionne un rappel ou un renvoi à un CR précédent ("comme décidé au CR08", "suite au CR13"), conserver cette référence dans la description.
- Un point sans action ni décision (simple constat positif, ex. "charpente posée, RAS") est valide : laisser action/responsable/delai à "".

RÈGLE ABSOLUE — ANTI-HALLUCINATION :
Tu ne dois JAMAIS inventer, deviner ou compléter une information absente de la transcription.
- Si un nom n'est pas prononcé → "nom": ""
- Si une entreprise n'est pas citée → "entreprise": ""
- Si un délai n'est pas précisé → "delai": ""
- Si la date de réunion n'est pas mentionnée → "date_reunion": null
- Si le lieu n'est pas mentionné → "lieu": null
- L'exemple ci-dessous est UNIQUEMENT pour illustrer le format JSON. Ne copie JAMAIS aucune valeur de cet exemple dans ta sortie.

Réponds UNIQUEMENT avec le JSON, sans texte avant ou après.

FORMAT DE SORTIE (utilise ces clés exactement, ne copie pas les valeurs) :
{
  "numero_cr": null,
  "date_reunion": "YYYY-MM-DD ou null",
  "lieu": "adresse exactement telle que mentionnée ou null",
  "presents": [{"nom": "prénom nom tel que prononcé", "qualite": "rôle tel que mentionné", "entreprise": "entreprise telle que mentionnée ou vide"}],
  "absents": [{"nom": "", "qualite": "", "entreprise": ""}],
  "lots": [
    {
      "numero": "01",
      "nom": "NOM DU LOT EN MAJUSCULES tel que mentionné",
      "entreprise": "entreprise telle que mentionnée ou vide",
      "points": [
        {
          "description": "description précise du point soulevé, avec dimensions/références si mentionnées",
          "decision": "décision prise si mentionnée, sinon vide",
          "action": "action à mener si mentionnée, sinon vide",
          "responsable": "nom ou entreprise responsable de l'action si mentionné, sinon vide",
          "delai": "délai en date absolue si mentionné, sinon vide"
        }
      ]
    }
  ],
  "divers": ["point divers tel que mentionné"],
  "prochaine_reunion": {"date": "YYYY-MM-DD ou vide", "lieu": "lieu tel que mentionné ou vide"},
  "diffusion": ["Nom — Entreprise (rôle) si mentionné"]
}"""


def _build_context_section(context_projet: dict | None) -> str:
    """Construit la section de contexte projet à injecter dans le prompt utilisateur."""
    if not context_projet:
        return ""
    lots = context_projet.get("lotsRecurrents") or []
    intervenants = context_projet.get("intervenants") or []
    if not lots and not intervenants:
        return ""
    lines = ["\nCONTEXTE DU PROJET (données des réunions précédentes) :"]
    if lots:
        lines.append("Lots habituels :")
        for lot in lots:
            entreprise = lot.get("entreprise", "")
            suffix = f" ({entreprise})" if entreprise else ""
            lines.append(f"  - LOT {lot.get('numero', '')} {lot.get('nom', '')}{suffix}")
    if intervenants:
        lines.append("Intervenants habituels :")
        for p in intervenants:
            parts = [p.get("nom", ""), p.get("qualite", ""), p.get("entreprise", "")]
            lines.append("  - " + " — ".join(x for x in parts if x))
    lines.append(
        "Utilise ces informations pour mieux identifier les intervenants si leur nom "
        "est partiellement prononcé ou mal transcrit. Tout nouvel intervenant mentionné "
        "dans la transcription doit être ajouté normalement."
    )
    return "\n".join(lines)


def structure_with_mistral(transcription: str, projet: str, context_projet: dict | None = None) -> dict:
    api_key = os.environ.get("MISTRAL_API_KEY")
    if not api_key:
        print("[ERREUR] MISTRAL_API_KEY non définie. Utilisez le mode mock.")
        sys.exit(1)

    try:
        from mistralai.client import Mistral
    except ImportError:
        print("[ERREUR] pip install mistralai")
        sys.exit(1)

    client = Mistral(api_key=api_key)
    print("Structuration via Mistral Medium 3...")

    context_section = _build_context_section(context_projet)

    user_prompt = f"""Voici la transcription d'une réunion de chantier.
Projet : {projet}{context_section}

Transcription :
{transcription}

Produis le compte rendu structuré en JSON."""

    response = client.chat.complete(
        model="mistral-medium-latest",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.1,
    )

    raw = response.choices[0].message.content.strip()
    # Nettoie les éventuels blocs markdown ```json ... ```
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1].rsplit("```", 1)[0]

    return json.loads(raw)


MOCK_CR = {
    "numero_cr": None,
    "date_reunion": "2025-05-28",
    "lieu": "12 rue des Lilas, Lyon",
    "presents": [
        {"nom": "Jean-Marc Aubert", "qualite": "Architecte", "entreprise": "Aubert et Associés"},
        {"nom": "Michel Renard", "qualite": "Responsable chantier", "entreprise": "Bâticorp"},
        {"nom": "Sophie Leclaire", "qualite": "Électricienne", "entreprise": "ElecPro"},
        {"nom": "Thomas Blanc", "qualite": "Menuisier", "entreprise": "Menuiserie Blanche"},
    ],
    "absents": [
        {"nom": "M. Fabre", "qualite": "Plombier", "entreprise": ""},
    ],
    "lots": [
        {
            "numero": "01",
            "nom": "GROS ŒUVRE",
            "entreprise": "Bâticorp",
            "points": [
                {
                    "description": "Coulage du dallage du garage",
                    "decision": "Coulage prévu le 01/06/2025 sous réserve météo",
                    "action": "Confirmer météo et procéder au coulage",
                    "responsable": "Michel Renard (Bâticorp)",
                    "delai": "01/06/2025",
                },
                {
                    "description": "Fissure de 2mm sur le mur de refend (salon)",
                    "decision": "Constat avec bureau de contrôle à prévoir",
                    "action": "Organiser constat avec bureau de contrôle",
                    "responsable": "Jean-Marc Aubert",
                    "delai": "Semaine prochaine",
                },
                {
                    "description": "PV de résistance du béton",
                    "decision": "",
                    "action": "Envoyer le PV de résistance du béton",
                    "responsable": "Michel Renard (Bâticorp)",
                    "delai": "Vendredi 30/05/2025",
                },
            ],
        },
        {
            "numero": "02",
            "nom": "CHARPENTE",
            "entreprise": "",
            "points": [
                {
                    "description": "Charpente posée — aucune remarque",
                    "decision": "",
                    "action": "",
                    "responsable": "",
                    "delai": "",
                }
            ],
        },
        {
            "numero": "03",
            "nom": "ÉLECTRICITÉ",
            "entreprise": "ElecPro",
            "points": [
                {
                    "description": "Passage des gaines dans cloisons R+1 : 80% terminé (reste chambre parentale et salle de bain)",
                    "decision": "Fin prévue fin de semaine prochaine",
                    "action": "Terminer passage des gaines",
                    "responsable": "Sophie Leclaire (ElecPro)",
                    "delai": "Fin de semaine du 02/06/2025",
                },
                {
                    "description": "Plans de cuisine nécessaires pour positionnement des prises",
                    "decision": "",
                    "action": "Envoyer plans de cuisine définitifs",
                    "responsable": "Jean-Marc Aubert",
                    "delai": "Ce soir 28/05/2025",
                },
            ],
        },
        {
            "numero": "04",
            "nom": "MENUISERIES EXTÉRIEURES",
            "entreprise": "Menuiserie Blanche",
            "points": [
                {
                    "description": "Fenêtre cuisine : problème de joint, fermeture défectueuse",
                    "decision": "Intervention fabricant nécessaire",
                    "action": "Contacter le fabricant pour intervention",
                    "responsable": "Thomas Blanc (Menuiserie Blanche)",
                    "delai": "Cette semaine",
                },
                {
                    "description": "Fenêtres R+1 non livrées — retard fournisseur",
                    "decision": "Nouveau délai annoncé : 15/06/2025",
                    "action": "Suivi livraison",
                    "responsable": "Thomas Blanc (Menuiserie Blanche)",
                    "delai": "15/06/2025",
                },
            ],
        },
    ],
    "divers": [
        "Sécurisation accès chantier : cadenas manquant sur le portail — Michel Renard s'en charge le 29/05/2025"
    ],
    "prochaine_reunion": {"date": "2025-06-10", "lieu": "Chantier, 12 rue des Lilas, Lyon"},
    "diffusion": [
        "Jean-Marc Aubert (Aubert et Associés)",
        "Michel Renard (Bâticorp)",
        "Sophie Leclaire (ElecPro)",
        "Thomas Blanc (Menuiserie Blanche)",
        "M. Fabre (Plombier)",
        "M. Dupont (Maître d'ouvrage)",
    ],
}


def structure_cr(transcription: str, projet: str, use_mock: bool, context_projet: dict | None = None) -> dict:
    if use_mock:
        print("[MOCK] Utilisation du CR de démonstration.")
        return MOCK_CR
    return structure_with_mistral(transcription, projet, context_projet)


# --------------------------------------------------------------------------- #
# 3. EXPORT WORD
# --------------------------------------------------------------------------- #

def _format_timestamp(iso_str: str) -> str:
    """Converts ISO timestamp to French format: DD/MM/YYYY HH:mm"""
    try:
        dt = datetime.fromisoformat(iso_str.replace("Z", "+00:00"))
        return dt.strftime("%d/%m/%Y %H:%M")
    except Exception:
        return iso_str


def normalize_cr(cr: dict) -> dict:
    """Normalise les champs qui peuvent être liste ou dict au lieu de str."""
    for lot in cr.get("lots", []):
        for point in lot.get("points", []):
            for key in ("description", "decision", "action", "responsable", "delai"):
                val = point.get(key, "")
                if isinstance(val, list):
                    parts = []
                    for item in val:
                        if isinstance(item, dict):
                            parts.append(
                                " — ".join(str(v) for v in item.values() if v)
                            )
                        else:
                            parts.append(str(item))
                    point[key] = " | ".join(parts)
                elif isinstance(val, dict):
                    point[key] = " — ".join(str(v) for v in val.values() if v)
                else:
                    point[key] = str(val) if val else ""
    return cr


def _add_cabinet_header_docx(doc, cabinet: dict | None):
    """Insère logo + infos cabinet en haut du document Word."""
    if not cabinet:
        return
    try:
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
    except ImportError:
        return

    has_logo = cabinet.get("logo") and cabinet.get("logoMime")
    has_info = any(cabinet.get(k) for k in ("nomCabinet", "adresse", "telephone", "email", "siteWeb"))

    if not has_logo and not has_info:
        return

    if has_logo:
        p = doc.add_paragraph()
        run = p.add_run()
        try:
            img_bytes = base64.b64decode(cabinet["logo"])
            run.add_picture(io.BytesIO(img_bytes), height=Inches(0.6))
        except Exception:
            pass
        p.paragraph_format.space_after = Pt(2)

    if has_info:
        info_parts = []
        if cabinet.get("nomCabinet"):
            info_parts.append(("bold", cabinet["nomCabinet"]))
        for key in ("adresse", "telephone", "email", "siteWeb"):
            if cabinet.get(key):
                info_parts.append(("normal", cabinet[key]))

        p = doc.add_paragraph()
        for i, (style_type, text) in enumerate(info_parts):
            if i > 0:
                p.add_run("  ·  ")
            run = p.add_run(text)
            run.font.size = Pt(8)
            if style_type == "bold":
                run.bold = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Ligne de séparation
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def export_word(cr: dict, projet: str, output_path: str, cabinet: dict | None = None):
    cr = normalize_cr(cr)
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[ERREUR] pip install python-docx")
        sys.exit(1)

    doc = Document()

    # Style de base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def heading(text, level=1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13 if level == 1 else 11)
        if level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_line(label, value):
        if not value:
            return
        p = doc.add_paragraph()
        run = p.add_run(f"{label} : ")
        run.bold = True
        p.add_run(str(value))
        p.paragraph_format.space_after = Pt(2)

    # Entête cabinet
    _add_cabinet_header_docx(doc, cabinet)

    # En-tête
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre.add_run("COMPTE RENDU DE CHANTIER")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    if cr.get("numero_cr"):
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(f"N° {cr['numero_cr']}").bold = True

    doc.add_paragraph()
    add_line("Opération", projet)
    add_line("Date", cr.get("date_reunion", ""))
    add_line("Lieu", cr.get("lieu", ""))
    doc.add_paragraph()

    # Présents / Absents
    heading("PRÉSENTS", 2)
    for p in cr.get("presents", []):
        parts = [p.get("nom", ""), p.get("qualite", ""), p.get("entreprise", "")]
        doc.add_paragraph("- " + " — ".join(x for x in parts if x), style="List Bullet")

    absents = cr.get("absents", [])
    if absents:
        heading("ABSENTS EXCUSÉS", 2)
        for p in absents:
            parts = [p.get("nom", ""), p.get("qualite", ""), p.get("entreprise", "")]
            doc.add_paragraph("- " + " — ".join(x for x in parts if x), style="List Bullet")

    doc.add_paragraph()

    # Lots
    heading("POINTS ABORDÉS PAR LOT")
    for lot in cr.get("lots", []):
        lot_title = f"LOT {lot.get('numero', '')} — {lot.get('nom', '')}"
        if lot.get("entreprise"):
            lot_title += f" ({lot['entreprise']})"
        heading(lot_title, 2)

        for point in lot.get("points", []):
            if point.get("description"):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(point["description"])

            for label, key in [("Décision", "decision"), ("Action", "action"),
                                ("Responsable", "responsable"), ("Délai", "delai")]:
                if point.get(key):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(24)
                    r = p.add_run(f"{label} : ")
                    r.bold = True
                    r.font.size = Pt(9)
                    p.add_run(point[key]).font.size = Pt(9)
                    p.paragraph_format.space_after = Pt(1)

            for photo in point.get("photos", []):
                try:
                    img_bytes = base64.b64decode(photo["data"])
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(24)
                    p.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(2.0))
                    if photo.get("timestamp"):
                        cap = doc.add_paragraph()
                        cap.paragraph_format.left_indent = Pt(24)
                        run2 = cap.add_run(_format_timestamp(photo["timestamp"]))
                        run2.italic = True
                        run2.font.size = Pt(8)
                        cap.paragraph_format.space_after = Pt(1)
                except Exception:
                    pass

        doc.add_paragraph()

    # Divers
    divers = [d for d in cr.get("divers", []) if d]
    if divers:
        heading("QUESTIONS / DIVERS", 2)
        for d in divers:
            doc.add_paragraph("- " + d, style="List Bullet")
        doc.add_paragraph()

    # Prochaine réunion
    pr = cr.get("prochaine_reunion", {})
    if pr.get("date") or pr.get("lieu"):
        heading("PROCHAINE RÉUNION", 2)
        add_line("Date", pr.get("date", ""))
        add_line("Lieu", pr.get("lieu", ""))
        doc.add_paragraph()

    # Diffusion
    diffusion = [d for d in cr.get("diffusion", []) if d]
    if diffusion:
        heading("DIFFUSION", 2)
        for d in diffusion:
            doc.add_paragraph("- " + d, style="List Bullet")

    doc.save(output_path)
    print(f"CR exporté : {output_path}")


# --------------------------------------------------------------------------- #
# 4. EXPORT PDF (fpdf2)
# --------------------------------------------------------------------------- #

def _pdf_safe(text: str) -> str:
    """Remplace les caractères hors Latin-1 courants dans les CR générés par Mistral."""
    replacements = {
        "•": "-",   # •  bullet
        "–": "-",   # –  en dash
        "—": "-",   # —  em dash
        "’": "'",   # '  right single quote
        "‘": "'",   # '  left single quote
        "“": '"',   # "  left double quote
        "”": '"',   # "  right double quote
        "…": "...", # …  ellipsis
        "Œ": "OE",  # Œ
        "œ": "oe",  # œ
        "æ": "ae",  # æ
        "Æ": "AE",  # Æ
    }
    for char, repl in replacements.items():
        text = text.replace(char, repl)
    # Encode en Latin-1, remplace ce qui reste
    return text.encode("latin-1", errors="replace").decode("latin-1")


def export_pdf(cr: dict, projet: str, output_path: str, cabinet: dict | None = None):
    """Génère un PDF du CR de chantier via fpdf2."""
    try:
        from fpdf import FPDF
    except ImportError:
        print("[ERREUR] pip install fpdf2")
        sys.exit(1)

    cr = normalize_cr(cr)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(20, 20, 20)
    pdf.add_page()

    BLUE = (31, 73, 125)
    GREY = (100, 100, 100)

    def set_blue():
        pdf.set_text_color(*BLUE)

    def set_black():
        pdf.set_text_color(0, 0, 0)

    # ── Entête cabinet ────────────────────────────────────────────────────────
    if cabinet:
        has_logo = cabinet.get("logo") and cabinet.get("logoMime")
        has_info = any(cabinet.get(k) for k in ("nomCabinet", "adresse", "telephone", "email", "siteWeb"))

        if has_logo:
            try:
                img_bytes = base64.b64decode(cabinet["logo"])
                pdf.image(io.BytesIO(img_bytes), x=20, y=pdf.get_y(), h=12, keep_aspect_ratio=True)
                pdf.set_y(pdf.get_y() + 14)
            except Exception:
                pass

        if has_info:
            info_parts = []
            if cabinet.get("nomCabinet"):
                info_parts.append(("B", cabinet["nomCabinet"]))
            for key in ("adresse", "telephone", "email", "siteWeb"):
                if cabinet.get(key):
                    info_parts.append(("", cabinet[key]))
            pdf.set_text_color(*GREY)
            for i, (style_type, text) in enumerate(info_parts):
                pdf.set_font("Helvetica", style_type, 8)
                if i > 0:
                    pdf.set_font("Helvetica", "", 8)
                    pdf.write(5, "  ·  ")
                    pdf.set_font("Helvetica", style_type, 8)
                pdf.write(5, text)
            pdf.ln(5)
            set_black()

        # Ligne de séparation
        pdf.set_draw_color(*BLUE)
        pdf.set_line_width(0.3)
        pdf.line(pdf.l_margin, pdf.get_y(), 210 - pdf.r_margin, pdf.get_y())
        pdf.ln(4)

    def section_title(text: str):
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 11)
        set_blue()
        pdf.cell(0, 8, text, new_x="LMARGIN", new_y="NEXT")
        set_black()
        pdf.set_draw_color(*BLUE)
        pdf.set_line_width(0.3)
        pdf.line(pdf.l_margin, pdf.get_y(), 210 - pdf.r_margin, pdf.get_y())
        pdf.ln(2)

    def info_line(label: str, value: str):
        if not value:
            return
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(40, 7, f"{label} :", new_x="RIGHT", new_y="TOP")
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 7, value, new_x="LMARGIN", new_y="NEXT")

    # ── Titre ──────────────────────────────────────────────────────────────
    pdf.set_font("Helvetica", "B", 16)
    set_blue()
    pdf.cell(0, 10, "COMPTE RENDU DE CHANTIER", align="C", new_x="LMARGIN", new_y="NEXT")
    if cr.get("numero_cr"):
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, f"N° {cr['numero_cr']}", align="C", new_x="LMARGIN", new_y="NEXT")
    set_black()
    pdf.set_draw_color(*BLUE)
    pdf.set_line_width(0.5)
    pdf.line(pdf.l_margin, pdf.get_y() + 2, 210 - pdf.r_margin, pdf.get_y() + 2)
    pdf.ln(6)

    # ── En-tête ─────────────────────────────────────────────────────────────
    info_line("Op\xe9ration", projet)
    info_line("Date", cr.get("date_reunion", ""))
    info_line("Lieu", cr.get("lieu", ""))
    pdf.ln(2)

    # ── Présents ─────────────────────────────────────────────────────────────
    presents = cr.get("presents", [])
    if presents:
        section_title("PR\xc9SENTS")
        pdf.set_font("Helvetica", "", 10)
        for p in presents:
            parts = [p.get("nom", ""), p.get("qualite", ""), p.get("entreprise", "")]
            line = " — ".join(x for x in parts if x)
            pdf.set_x(pdf.l_margin + 3)
            pdf.cell(5, 6, "•", new_x="RIGHT", new_y="TOP")
            pdf.multi_cell(0, 6, line, new_x="LMARGIN", new_y="NEXT")

    absents = cr.get("absents", [])
    if absents:
        pdf.ln(1)
        pdf.set_font("Helvetica", "BI", 9)
        pdf.set_text_color(*GREY)
        pdf.cell(0, 6, "Absents excus\xe9s", new_x="LMARGIN", new_y="NEXT")
        set_black()
        pdf.set_font("Helvetica", "", 10)
        for p in absents:
            parts = [p.get("nom", ""), p.get("qualite", ""), p.get("entreprise", "")]
            line = " — ".join(x for x in parts if x)
            pdf.set_x(pdf.l_margin + 3)
            pdf.cell(5, 6, "•", new_x="RIGHT", new_y="TOP")
            pdf.multi_cell(0, 6, line, new_x="LMARGIN", new_y="NEXT")

    # ── Lots ─────────────────────────────────────────────────────────────────
    section_title("POINTS ABORD\xc9S PAR LOT")

    for lot in cr.get("lots", []):
        lot_title = f"LOT {lot.get('numero', '')} — {lot.get('nom', '')}"
        if lot.get("entreprise"):
            lot_title += f" ({lot['entreprise']})"
        pdf.set_font("Helvetica", "B", 10)
        set_blue()
        pdf.cell(0, 7, lot_title, new_x="LMARGIN", new_y="NEXT")
        set_black()

        for point in lot.get("points", []):
            desc = point.get("description", "")
            if desc:
                pdf.set_font("Helvetica", "", 10)
                pdf.set_x(pdf.l_margin + 5)
                pdf.cell(5, 6, "•", new_x="RIGHT", new_y="TOP")
                pdf.multi_cell(0, 6, desc, new_x="LMARGIN", new_y="NEXT")

            for label, key in [("D\xe9cision", "decision"), ("Action", "action"),
                                ("Responsable", "responsable"), ("D\xe9lai", "delai")]:
                val = point.get(key, "")
                if val:
                    pdf.set_x(pdf.l_margin + 10)
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.cell(28, 5, f"{label} :", new_x="RIGHT", new_y="TOP")
                    pdf.set_font("Helvetica", "", 9)
                    pdf.multi_cell(0, 5, val, new_x="LMARGIN", new_y="NEXT")

            for photo in point.get("photos", []):
                try:
                    img_bytes = base64.b64decode(photo.get("data", ""))
                    y_start = pdf.get_y()
                    pdf.image(
                        io.BytesIO(img_bytes),
                        x=pdf.l_margin + 10,
                        y=y_start,
                        w=50,
                        h=35,
                        keep_aspect_ratio=True,
                    )
                    pdf.set_y(y_start + 37)
                    if photo.get("timestamp"):
                        pdf.set_x(pdf.l_margin + 10)
                        pdf.set_font("Helvetica", "I", 7)
                        pdf.set_text_color(*GREY)
                        pdf.cell(50, 4, _format_timestamp(photo["timestamp"]),
                                 align="C", new_x="LMARGIN", new_y="NEXT")
                        set_black()
                    pdf.ln(2)
                except Exception:
                    pass

            pdf.ln(1)
        pdf.ln(2)

    # ── Divers ────────────────────────────────────────────────────────────────
    divers = [d for d in cr.get("divers", []) if d]
    if divers:
        section_title("QUESTIONS / DIVERS")
        pdf.set_font("Helvetica", "", 10)
        for d in divers:
            pdf.set_x(pdf.l_margin + 3)
            pdf.cell(5, 6, "•", new_x="RIGHT", new_y="TOP")
            pdf.multi_cell(0, 6, d, new_x="LMARGIN", new_y="NEXT")

    # ── Prochaine réunion ─────────────────────────────────────────────────────
    pr = cr.get("prochaine_reunion", {})
    if pr.get("date") or pr.get("lieu"):
        section_title("PROCHAINE R\xc9UNION")
        info_line("Date", pr.get("date", ""))
        info_line("Lieu", pr.get("lieu", ""))

    # ── Diffusion ─────────────────────────────────────────────────────────────
    diffusion = [d for d in cr.get("diffusion", []) if d]
    if diffusion:
        section_title("DIFFUSION")
        pdf.set_font("Helvetica", "", 10)
        for d in diffusion:
            pdf.set_x(pdf.l_margin + 3)
            pdf.cell(5, 6, "•", new_x="RIGHT", new_y="TOP")
            pdf.multi_cell(0, 6, d, new_x="LMARGIN", new_y="NEXT")

    pdf.output(output_path)
    print(f"CR PDF export\xe9 : {output_path}")


# --------------------------------------------------------------------------- #
# 5. MAIN
# --------------------------------------------------------------------------- #

def main():
    parser = argparse.ArgumentParser(description="Pipeline CR de chantier")
    parser.add_argument("--audio", help="Chemin vers le fichier audio (MP3, M4A, WAV)")
    parser.add_argument("--transcription", help="Chemin vers un fichier texte de transcription")
    parser.add_argument("--projet", default="Maison Dupont", help="Nom du projet")
    parser.add_argument("--output", default="compte_rendu.docx", help="Fichier Word de sortie")
    args = parser.parse_args()

    use_mock = not args.audio and not args.transcription and not os.environ.get("GLADIA_API_KEY")

    print("=== Pipeline CR de Chantier ===")
    print(f"Projet : {args.projet}")
    print(f"Mode   : {'MOCK (démonstration)' if use_mock else 'RÉEL'}")
    print()

    transcription = get_transcription(args)

    transcription_path = Path(args.output).stem + "_transcription.txt"
    Path(transcription_path).write_text(transcription, encoding="utf-8")
    print(f"Transcription sauvegardée : {transcription_path}")

    cr = structure_cr(transcription, args.projet, use_mock)

    print()
    print("Structure JSON générée :")
    print(json.dumps(cr, ensure_ascii=False, indent=2))
    print()

    export_word(cr, args.projet, args.output)
    print(f"\nTerminé. Ouvrez {args.output}")


if __name__ == "__main__":
    main()
